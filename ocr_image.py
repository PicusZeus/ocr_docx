#!/usr/bin/python3

import concurrent.futures
from PIL import Image
import os
import re
import pickle
import docx
import logging
from pathlib import Path
import pytesseract
from nltk.tokenize.punkt import PunktLanguageVars
from string import punctuation
from multiprocessing import Lock


lemmatizer = PunktLanguageVars()

logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)
format_log = logging.Formatter('%(levelname)s:%(name)s:%(message)s')
log_path = Path(Path.cwd() / 'ocr.log')
fileHandler = logging.FileHandler(log_path)
fileHandler.setFormatter(format_log)
logger.addHandler(fileHandler)
stream = logging.StreamHandler()
stream.setFormatter(format_log)
logger.addHandler(stream)

lock = Lock()
new_paragraphs = []

paragraphs = []
doc = docx.Document()
ocr_texts = []


def ocr_text_to_docx(file, image_folder='images', output_file='ocred.docx', lang='grc'):

    if lang == 'grc':
        global greek_voc
        greek_voc = pickle.load(open('greek_voc.pickle', 'rb'))

    # make a list of paragraphs in the doc that contain images
    if not Path(file).is_file():
        logger.error("No such file")
        raise FileNotFoundError
    image_paragraphs = []
    doc = docx.Document(file)
    for par in doc.paragraphs:
        if 'graphicData' in par._p.xml:
            image_paragraphs.append(par)

    # list of extracted images from doc file
    images = os.listdir(image_folder)

    # it has to be sorted, images are written in a pattern "image1"

    # filter out only relevant files
    images = list(filter(re.compile(r'.*\.(png|jpg|JPG|jpeg|wmf)').match, images))

    # sort them in ascending order
    images = sorted(images)
    # so that order be image9, image10, and not image1, image10
    images.sort(key=lambda x:len(x))

    with concurrent.futures.ProcessPoolExecutor() as executor:
        res = [executor.submit(ocr_image, index, image_folder, images, lang) for index in range(len(image_paragraphs))]
        for f in concurrent.futures.as_completed(res):
            index, checked_text = f.result()
            par = image_paragraphs[index]
            added_paragraph = par.insert_paragraph_before()
            for w in checked_text:
                if w in punctuation:
                    added_paragraph.add_run(text=w)
                else:
                    ws = ' ' + w
                    if '$' in w:

                        ws = ws.replace('$', '')
                        added_paragraph.add_run(text=ws).bold = True
                    else:
                        added_paragraph.add_run(text=ws)
            logger.info(f'ocred {index + 1} images')

    doc.save(output_file)


def ocr_image(index, image_folder, images, lang):

    with lock:
        file_path = (Path.cwd() / image_folder / images[index])
        logger.info(f'processing {index} image')

        text = pytesseract.image_to_string(Image.open(str(file_path)), lang=lang)

        text = text.replace('-\n', '')
        text = text.replace('\n', ' ')

        tokenized_text = lemmatizer.word_tokenize(text)

        checked_text = []

        for w in tokenized_text:

            if lang == 'grc' and w.lower() not in greek_voc:
                checked_text.append(w + "$")
            else:
                checked_text.append(w)

    return [index, checked_text]


